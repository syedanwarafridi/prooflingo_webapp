# import os
# import xml.etree.ElementTree as ET
# from docx import Document

# def create_directories():
#     if not os.path.exists('./source_files'):
#         os.makedirs('./source_files')
#     if not os.path.exists('./target_files'):
#         os.makedirs('./target_files')

# # Trados --> .sdlxliff, .txlf
# def billingual_sengmentation_trados(file_path):
#     create_directories()
#     tree = ET.parse(file_path)
#     root = tree.getroot()
#     source_text = []
#     target_text = []

#     for trans_unit in root.findall('.//{urn:oasis:names:tc:xliff:document:1.2}trans-unit'):
#         source = trans_unit.find('{urn:oasis:names:tc:xliff:document:1.2}source')
#         target = trans_unit.find('{urn:oasis:names:tc:xliff:document:1.2}target')
        
#         if source is not None:
#             source_text.append(source.text)
        
#         if target is not None:
#             target_text.append(target.text)

#     source_doc = Document()
#     source_doc.add_heading('Source Text', level=1)
#     for text in source_text:
#         source_doc.add_paragraph(text)
#     source_path = f'./source_files/{os.path.basename(file_path).split(".")[0]}_source.docx'
#     source_doc.save(source_path)

#     target_doc = Document()
#     target_doc.add_heading('Target Text', level=1)
#     for text in target_text:
#         target_doc.add_paragraph(text)
#     target_path = f'./target_files/{os.path.basename(file_path).split(".")[0]}_target.docx'
#     target_doc.save(target_path)

#     # print("Conversion to DOCX successful.")
#     return source_path, target_path

# # XML --> .xml
# def billingual_segmentation_xml(file_path):
#     create_directories()
#     tree = ET.parse(file_path)
#     root = tree.getroot()
#     source_text = []
#     target_text = []

#     for segment in root.findall('segment'):
#         source = segment.find('source')
#         target = segment.find('target')
        
#         if source is not None:
#             source_text.append(source.text)
        
#         if target is not None:
#             target_text.append(target.text)

#     source_doc = Document()
#     source_doc.add_heading('Source Text', level=1)
#     for text in source_text:
#         source_doc.add_paragraph(text)
#     source_path = f'./source_files/{os.path.basename(file_path).split(".")[0]}_source.docx'
#     source_doc.save(source_path)

#     target_doc = Document()
#     target_doc.add_heading('Target Text', level=1)
#     for text in target_text:
#         target_doc.add_paragraph(text)
#     target_path = f'./target_files/{os.path.basename(file_path).split(".")[0]}_target.docx'
#     target_doc.save(target_path)

#     # print("Conversion to DOCX successful.")
#     return source_path, target_path

# # MemoQ --> .mqxliff
# def billingual_segmentation_mqxliff(file_path):
#     create_directories()
#     tree = ET.parse(file_path)
#     root = tree.getroot()
#     source_text = []
#     target_text = []

#     for trans_unit in root.findall('.//{urn:oasis:names:tc:xliff:document:2.0}unit'):
#         source = trans_unit.find('{urn:oasis:names:tc:xliff:document:2.0}segment/{urn:oasis:names:tc:xliff:document:2.0}source')
#         target = trans_unit.find('{urn:oasis:names:tc:xliff:document:2.0}segment/{urn:oasis:names:tc:xliff:document:2.0}target')
        
#         if source is not None:
#             source_text.append(source.text)
        
#         if target is not None:
#             target_text.append(target.text)

#     source_doc = Document()
#     source_doc.add_heading('Source Text', level=1)
#     for text in source_text:
#         source_doc.add_paragraph(text)
#     source_path = f'./source_files/{os.path.basename(file_path).split(".")[0]}_source.docx'
#     source_doc.save(source_path)

#     target_doc = Document()
#     target_doc.add_heading('Target Text', level=1)
#     for text in target_text:
#         target_doc.add_paragraph(text)
#     target_path = f'./target_files/{os.path.basename(file_path).split(".")[0]}_target.docx'
#     target_doc.save(target_path)

#     # print("Conversion to DOCX successful.")
#     return source_path, target_path

# # Memsource --> .mxliff
# def billingual_segmentation_mxliff(file_path):
#     create_directories()
#     tree = ET.parse(file_path)
#     root = tree.getroot()
#     source_text = []
#     target_text = []

#     for unit in root.findall('.//{urn:oasis:names:tc:xliff:document:2.0}unit'):
#         source = unit.find('{urn:oasis:names:tc:xliff:document:2.0}segment/{urn:oasis:names:tc:xliff:document:2.0}source')
#         target = unit.find('{urn:oasis:names:tc:xliff:document:2.0}segment/{urn:oasis:names:tc:xliff:document:2.0}target')
        
#         if source is not None:
#             source_text.append(source.text)
        
#         if target is not None:
#             target_text.append(target.text)

#     source_doc = Document()
#     source_doc.add_heading('Source Text', level=1)
#     for text in source_text:
#         source_doc.add_paragraph(text)
#     source_path = f'./source_files/{os.path.basename(file_path).split(".")[0]}_source.docx'
#     source_doc.save(source_path)

#     target_doc = Document()
#     target_doc.add_heading('Target Text', level=1)
#     for text in target_text:
#         target_doc.add_paragraph(text)
#     target_path = f'./target_files/{os.path.basename(file_path).split(".")[0]}_target.docx'
#     target_doc.save(target_path)

#     # print("Conversion to DOCX successful.")
#     return source_path, target_path

# # SDLPPX --> .sdlppx
# def billingual_segmentation_sdlppx(file_path):
#     create_directories()
#     tree = ET.parse(file_path)
#     root = tree.getroot()
#     source_doc = Document()
#     source_doc.add_heading('Source Text', level=1)
#     target_doc = Document()
#     target_doc.add_heading('Target Text', level=1)
    
#     for segment in root.findall(".//Segment"):
#         source_text = segment.find("Source").text
#         target_text = segment.find("Target").text

#         source_doc.add_paragraph(source_text)
#         target_doc.add_paragraph(target_text)

#     source_path = f'./source_files/{os.path.basename(file_path).split(".")[0]}_source.docx'
#     source_doc.save(source_path)
#     target_path = f'./target_files/{os.path.basename(file_path).split(".")[0]}_target.docx'
#     target_doc.save(target_path)

#     # print("Documents created successfully.")
#     return source_path, target_path

# # XLIFF --> .xliff
# def billingual_segmentation_xliff(file_path):
#     create_directories()
#     tree = ET.parse(file_path)
#     root = tree.getroot()
#     source_text = []
#     target_text = []

#     for trans_unit in root.findall('.//{urn:oasis:names:tc:xliff:document:1.2}trans-unit'):
#         source = trans_unit.find('{urn:oasis:names:tc:xliff:document:1.2}source')
#         target = trans_unit.find('{urn:oasis:names:tc:xliff:document:1.2}target')
        
#         if source is not None:
#             source_text.append(source.text)
        
#         if target is not None:
#             target_text.append(target.text)

#     source_doc = Document()
#     source_doc.add_heading('Source Text', level=1)
#     for text in source_text:
#         source_doc.add_paragraph(text)
#     source_path = f'./source_files/{os.path.basename(file_path).split(".")[0]}_source.docx'
#     source_doc.save(source_path)

#     target_doc = Document()
#     target_doc.add_heading('Target Text', level=1)
#     for text in target_text:
#         target_doc.add_paragraph(text)
#     target_path = f'./target_files/{os.path.basename(file_path).split(".")[0]}_target.docx'
#     target_doc.save(target_path)

#     # print("Conversion to DOCX successful.")
#     return source_path, target_path

# # Example usage:
# # source_path, target_path = billingual_sengmentation_trados('trados.sdlxliff')
# # source_path, target_path = billingual_segmentation_xml('example.xml')
# # source_path, target_path = billingual_segmentation_mqxliff('memoq.mqxliff')
# # source_path, target_path = billingual_segmentation_mxliff('memsource.mxliff')
# # source_path, target_path = billingual_segmentation_sdlppx('hello.sdlppx')
# # source_path, target_path = billingual_segmentation_xliff('sample.xliff')


import os
import xml.etree.ElementTree as ET
from docx import Document
from io import BytesIO
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.core.files.uploadedfile import InMemoryUploadedFile

def create_directories():
    if not os.path.exists('./source_files'):
        os.makedirs('./source_files')
    if not os.path.exists('./target_files'):
        os.makedirs('./target_files')

def create_docx_file(texts, heading):
    doc = Document()
    doc.add_heading(heading, level=1)
    for text in texts:
        doc.add_paragraph(text)
    return doc

def save_docx_to_bytes(doc):
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# Trados --> .sdlxliff, .txlf
def billingual_sengmentation_trados(file_path):
    tree = ET.parse(file_path)
    root = tree.getroot()
    source_text = []
    target_text = []

    for trans_unit in root.findall('.//{urn:oasis:names:tc:xliff:document:1.2}trans-unit'):
        source = trans_unit.find('{urn:oasis:names:tc:xliff:document:1.2}source')
        target = trans_unit.find('{urn:oasis:names:tc:xliff:document:1.2}target')
        
        if source is not None:
            source_text.append(source.text)
        
        if target is not None:
            target_text.append(target.text)

    source_doc = create_docx_file(source_text, 'Source Text')
    target_doc = create_docx_file(target_text, 'Target Text')

    return save_docx_to_bytes(source_doc), save_docx_to_bytes(target_doc)

# XML --> .xml
def billingual_segmentation_xml(file_path):
    tree = ET.parse(file_path)
    root = tree.getroot()
    source_text = []
    target_text = []

    for segment in root.findall('segment'):
        source = segment.find('source')
        target = segment.find('target')
        
        if source is not None:
            source_text.append(source.text)
        
        if target is not None:
            target_text.append(target.text)

    source_doc = create_docx_file(source_text, 'Source Text')
    target_doc = create_docx_file(target_text, 'Target Text')

    return save_docx_to_bytes(source_doc), save_docx_to_bytes(target_doc)

# MemoQ --> .mqxliff
def billingual_segmentation_mqxliff(file_path):
    tree = ET.parse(file_path)
    root = tree.getroot()
    source_text = []
    target_text = []

    for trans_unit in root.findall('.//{urn:oasis:names:tc:xliff:document:2.0}unit'):
        source = trans_unit.find('{urn:oasis:names:tc:xliff:document:2.0}segment/{urn:oasis:names:tc:xliff:document:2.0}source')
        target = trans_unit.find('{urn:oasis:names:tc:xliff:document:2.0}segment/{urn:oasis:names:tc:xliff:document:2.0}target')
        
        if source is not None:
            source_text.append(source.text)
        
        if target is not None:
            target_text.append(target.text)

    source_doc = create_docx_file(source_text, 'Source Text')
    target_doc = create_docx_file(target_text, 'Target Text')

    return save_docx_to_bytes(source_doc), save_docx_to_bytes(target_doc)

# Memsource --> .mxliff
def billingual_segmentation_mxliff(file_path):
    tree = ET.parse(file_path)
    root = tree.getroot()
    source_text = []
    target_text = []

    for unit in root.findall('.//{urn:oasis:names:tc:xliff:document:2.0}unit'):
        source = unit.find('{urn:oasis:names:tc:xliff:document:2.0}segment/{urn:oasis:names:tc:xliff:document:2.0}source')
        target = unit.find('{urn:oasis:names:tc:xliff:document:2.0}segment/{urn:oasis:names:tc:xliff:document:2.0}target')
        
        if source is not None:
            source_text.append(source.text)
        
        if target is not None:
            target_text.append(target.text)

    source_doc = create_docx_file(source_text, 'Source Text')
    target_doc = create_docx_file(target_text, 'Target Text')

    return save_docx_to_bytes(source_doc), save_docx_to_bytes(target_doc)

# SDLPPX --> .sdlppx
def billingual_segmentation_sdlppx(file_path):
    tree = ET.parse(file_path)
    root = tree.getroot()
    source_doc = Document()
    source_doc.add_heading('Source Text', level=1)
    target_doc = Document()
    target_doc.add_heading('Target Text', level=1)
    
    for segment in root.findall(".//Segment"):
        source_text = segment.find("Source").text
        target_text = segment.find("Target").text

        source_doc.add_paragraph(source_text)
        target_doc.add_paragraph(target_text)

    return save_docx_to_bytes(source_doc), save_docx_to_bytes(target_doc)

# XLIFF --> .xliff
def billingual_segmentation_xliff(file_path):
    tree = ET.parse(file_path)
    root = tree.getroot()
    source_text = []
    target_text = []

    for trans_unit in root.findall('.//{urn:oasis:names:tc:xliff:document:1.2}trans-unit'):
        source = trans_unit.find('{urn:oasis:names:tc:xliff:document:1.2}source')
        target = trans_unit.find('{urn:oasis:names:tc:xliff:document:1.2}target')
        
        if source is not None:
            source_text.append(source.text)
        
        if target is not None:
            target_text.append(target.text)

    source_doc = create_docx_file(source_text, 'Source Text')
    target_doc = create_docx_file(target_text, 'Target Text')

    return save_docx_to_bytes(source_doc), save_docx_to_bytes(target_doc)

